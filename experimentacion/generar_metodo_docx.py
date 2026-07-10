# -*- coding: utf-8 -*-
"""Metodo de priorizacion en DOCX (diseno simple)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/kpmguser/TesisNicoLLM/experimentacion/Metodo_de_Priorizacion.docx"
INK = RGBColor(0x0e, 0x1c, 0x30)
TEAL = RGBColor(0x0f, 0x76, 0x6b)

d = Document()
st = d.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)

def h1(t):
    p = d.add_heading(level=0)  # title
    r = p.add_run(t); r.font.color.rgb = INK; r.font.size = Pt(20)
def h2(t):
    p = d.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = INK
def h3(t):
    p = d.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = TEAL; r.font.size = Pt(11)
def para(runs, align=None):
    """runs: lista de (texto, bold)."""
    p = d.add_paragraph()
    if align: p.alignment = align
    for txt, b in runs:
        r = p.add_run(txt); r.bold = b
    return p
def txt(s):
    """Convierte '<b>..</b>' en runs."""
    out = []
    for i, part in enumerate(s.split("**")):
        if part: out.append((part, i % 2 == 1))
    return out

def table(rows, widths=None):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j); c.text = ""
            p = c.paragraphs[0]
            for tt, b in txt(cell):
                r = p.add_run(tt); r.bold = b or (i == 0)
                r.font.size = Pt(9)
                if i == 0: r.font.color.rgb = INK
    return t

# ---- Titulo ----
h1("Método de priorización de requisitos")
para(txt("Detalle del cálculo para ajustar los instrumentos de evaluación (cuestionarios y entrevistas). "
         "Rasgo distintivo: la priorización **integra el análisis ético** del requisito y el cálculo es "
         "**determinista y auditable** (sin IA en la fórmula: reproducible y explicable)."))
para(txt("Tesis · Nicolás Pérez — Fases 4–6, con enlace a Fases 2–3, 7 y 8."))

# 1
h2("1 · Idea general")
para(txt("Cada requisito se puntúa según cuánto **aporta** (beneficios y valores éticos) y cuánto **cuesta o "
         "arriesga** (costos y riesgos éticos). El puntaje resulta de una fórmula ponderada simple; los "
         "requisitos se ordenan de mayor a menor. Las dimensiones éticas no se inventan al priorizar: "
         "provienen del **análisis ético asistido por IA** (Fases 2–3), de modo que un requisito muy útil pero "
         "éticamente problemático **baja** en el ranking frente a otro de utilidad similar y menor riesgo."))

# 2
h2("2 · Piezas del método y escalas")
para(txt("Una **dimensión** es un criterio de valoración con un **tipo** (define su signo) y un **peso** "
         "(importancia). Cada dimensión se cruza con cada requisito mediante una **fuerza** (intensidad con que "
         "el requisito expresa esa dimensión)."))
table([
    ["Concepto", "Escala", "Significado"],
    ["Peso de la dimensión", "1 a 5", "Cuánto importa el criterio (lo fija el equipo; la IA lo sugiere para las éticas)."],
    ["Fuerza requisito × dimensión", "0 a 5", "Cuán intensamente el requisito activa la dimensión. **0 = no aplica.**"],
])
h3("Los cuatro tipos de dimensión")
table([
    ["Tipo", "Signo", "Ejemplos"],
    ["Beneficio", "+ suma", "Valor para el usuario, viabilidad técnica"],
    ["Valor ético", "+ suma", "Equidad y no discriminación, transparencia"],
    ["Costo", "− resta", "Costo de desarrollo, esfuerzo operativo"],
    ["Riesgo ético", "− resta", "Riesgo de privacidad, sesgo, opacidad"],
])

# 3
h2("3 · Los tres pasos del método")
for a, b in [
    ("Fase 4 — Definir dimensiones", "se declaran las dimensiones con tipo y peso (1–5). Las generales "
     "(beneficio, costo) aplican a todos los requisitos; las éticas (valor/riesgo) suelen ser restringidas: "
     "aplican solo donde el análisis las detectó."),
    ("Fase 5 — Evaluar la matriz", "para cada requisito vigente y dimensión aplicable se asigna la fuerza (0–5), "
     "con justificación. Donde no aplica, la fuerza es 0."),
    ("Fase 6 — Calcular el ranking", "se aplica la fórmula a cada requisito y se ordena de mayor a menor. "
     "Función determinista y auditable: mismos datos → mismo ranking."),
]:
    p = d.add_paragraph(style="List Bullet")
    r = p.add_run(a + ": "); r.bold = True
    p.add_run(b)

h3("Fórmula")
para([("Puntaje final = Σ(peso × fuerza) beneficio + Σ(peso × fuerza) valor ético "
       "− Σ(peso × fuerza) costo − Σ(peso × fuerza) riesgo ético", True)])
para(txt("Solo entran los requisitos **vigentes** (no eliminados ni reemplazados por una reformulación). El "
         "puntaje puede ser **negativo** si costos y riesgos superan a beneficios y valores. Los empates se "
         "resuelven por código (orden estable). El resultado incluye el **desglose** por las cuatro categorías, "
         "así siempre se explica por qué un requisito quedó donde quedó."))

# 4
h2("4 · Ejemplo numérico")
para(txt("Proyecto con cuatro dimensiones (peso entre paréntesis) y dos requisitos:"))
table([
    ["Requisito", "Valor usuario benef. (5)", "Equidad ético (4)", "Costo costo (3)", "Privacidad riesgo (4)", "Puntaje"],
    ["R-A · útil pero riesgoso", "5", "2", "3", "5", "4"],
    ["R-B · sólido y ético", "4", "5", "2", "1", "30"],
])
para(txt("R-A = (5×5) + (4×2) − (3×3) − (4×5) = 25 + 8 − 9 − 20 = **4**."))
para(txt("R-B = (5×4) + (4×5) − (3×2) − (4×1) = 20 + 20 − 6 − 4 = **30**."))
para(txt("Aunque R-A aporta más valor de usuario, su alto riesgo de privacidad y baja equidad lo hunden; R-B "
         "queda primero. Que la ética mueva el ranking es justo lo que conviene evaluar con los usuarios."))

# 5
h2("5 · Cómo se conecta con el resto del método")
for a, b in [
    ("De dónde salen las dimensiones éticas (Fases 2–3)", "el análisis asistido por IA detecta los temas "
     "éticos (actor afectado, tipo de daño, norma tensionada con citas reales) y propone dimensiones de valor "
     "o riesgo ético con peso sugerido; el humano las revisa, edita o descarta antes de priorizar."),
    ("Regla de arrastre (Fase 7)", "si un requisito derivado obligatorio (p. ej. una mitigación) queda con "
     "puntaje menor que el requisito del que proviene, se marca una alerta informativa de posible "
     "inconsistencia de prioridades."),
    ("Bandera ética y tablero (Fase 8)", "el tablero muestra ranking, desglose y una bandera ética por "
     "requisito derivada del análisis; el proyecto se puede exportar (JSON/CSV) para auditar."),
]:
    p = d.add_paragraph(style="List Bullet")
    r = p.add_run(a + ": "); r.bold = True
    p.add_run(b)

# 6
h2("6 · Constructos sugeridos para los instrumentos")
para(txt("Focos posibles para cuestionarios (Likert 1–5) y entrevistas, alineados con el método:"))
table([
    ["Constructo", "Qué indagar"],
    ["Utilidad percibida", "¿El ranking ayudó a decidir qué requisitos priorizar?"],
    ["Detección de riesgos", "¿La priorización visibilizó tensiones éticas no evidentes a simple vista?"],
    ["Transparencia / explicabilidad", "¿El desglose y la fórmula hicieron entendible el porqué del orden?"],
    ["Confianza", "¿Confiarían en este puntaje para justificar una decisión ante terceros?"],
    ["Control humano", "¿Fue adecuado poder editar pesos, fuerzas y las dimensiones propuestas por la IA?"],
    ["Carga / esfuerzo", "¿Cuánto costó completar la matriz? ¿Fue proporcional al valor obtenido?"],
    ["Adopción", "¿Usarían este método en un proyecto real? ¿En qué contexto sí o no?"],
])

d.save(OUT)
print("OK ->", OUT)
